import logging
from typing import Dict, Callable
import os
import PyPDF2
import docx
import pandas as pd
import json

from dataset_miner.file_types import EXTENSION_TO_LANGUAGE, get_file_extension

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Custom exception for text extraction errors."""

    pass


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF files."""
    logger.info(f"Extracting text from PDF: {file_path}")
    try:
        with open(file_path, "rb") as pdf_file_obj:
            pdf_reader = PyPDF2.PdfReader(pdf_file_obj)
            text = ""
            total_pages = len(pdf_reader.pages)
            logger.info(f"Processing {total_pages} pages from PDF")

            for i, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                    logger.debug(f"Processed page {i}/{total_pages}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {i}: {str(e)}")

            logger.info(f"Extracted {len(text)} characters from {file_path}")
            return text.strip()
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}") from e


def extract_text_from_txt(file_path: str, encoding: str = "utf-8") -> str:
    """Extract text from TXT files with flexible encoding support."""
    logger.info(f"Extracting text from TXT: {file_path}")
    encodings = [encoding, "utf-8", "latin-1", "ascii"]

    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as txt_file:
                text = txt_file.read()
            logger.info(f"Successfully read file with {enc} encoding")
            return text.strip()
        except UnicodeDecodeError:
            continue

    raise TextExtractionError(
        f"Failed to decode file with any of the attempted encodings: {encodings}"
    )


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files including tables and headers."""
    logger.info(f"Extracting text from DOCX: {file_path}")
    try:
        doc = docx.Document(file_path)
        content_parts = []

        # Extract headers
        for section in doc.sections:
            header = section.header
            if header.paragraphs:
                content_parts.append("=== Header ===")
                content_parts.extend(
                    p.text for p in header.paragraphs if p.text.strip()
                )

        # Extract main content
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content_parts.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            content_parts.append("=== Table ===")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    content_parts.append(row_text)

        text = "\n".join(content_parts)
        logger.info(f"Extracted {len(text)} characters from {file_path}")
        return text
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}") from e


def extract_text_from_json(file_path: str) -> str:
    """Extract text from JSON files with pretty formatting."""
    logger.info(f"Extracting text from JSON: {file_path}")
    try:
        with open(file_path, "r", encoding="utf-8") as json_file:
            data = json.load(json_file)
        text = json.dumps(data, indent=2, ensure_ascii=False)
        logger.info(f"Extracted {len(text)} characters from {file_path}")
        return text
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from JSON: {str(e)}") from e


def extract_text_from_csv(file_path: str) -> str:
    """Extract text from CSV files with advanced options."""
    logger.info(f"Extracting text from CSV: {file_path}")
    try:
        # Try different encodings
        encodings = ["utf-8", "latin-1", "ascii"]
        df = None

        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                break
            except UnicodeDecodeError:
                continue

        if df is None:
            raise TextExtractionError("Failed to read CSV with any supported encoding")

        # Convert to string with formatting
        text = df.to_string(index=False, na_rep="N/A")
        logger.info(f"Extracted {len(text)} characters from {file_path}")
        return text
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from CSV: {str(e)}") from e


def extract_text_from_xlsx(file_path: str) -> str:
    """Extract text from Excel files with sheet handling."""
    logger.info(f"Extracting text from Excel: {file_path}")
    try:
        sheets_dict = pd.read_excel(file_path, sheet_name=None)
        text_parts = []

        for sheet_name, sheet_data in sheets_dict.items():
            text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
            text_parts.append(sheet_data.to_string(index=False, na_rep="N/A"))
            text_parts.append("\n")

        text = "\n".join(text_parts)
        logger.info(f"Extracted {len(text)} characters from {file_path}")
        return text
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from Excel: {str(e)}") from e


def extract_text_from_code(file_path: str) -> str:
    """Extract text from code files with basic formatting."""
    logger.info(f"Extracting text from code file: {file_path}")
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="latin-1") as file:
                return file.read()
        except Exception as e:
            raise TextExtractionError(
                f"Failed to extract text from code file: {str(e)}"
            ) from e


# Mapping of file extensions to their extraction functions
EXTRACTION_FUNCTIONS: Dict[str, Callable[[str], str]] = {
    ".pdf": extract_text_from_pdf,
    ".txt": extract_text_from_txt,
    ".docx": extract_text_from_docx,
    ".json": extract_text_from_json,
    ".csv": extract_text_from_csv,
    ".xlsx": extract_text_from_xlsx,
    ".xls": extract_text_from_xlsx,
}


def extract_text(file_path: str) -> str:
    """
    Extract text from a file based on its extension.

    Args:
        file_path: Path to the file to extract text from

    Returns:
        Extracted text as a string

    Raises:
        TextExtractionError: If text extraction fails
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_extension = get_file_extension(file_path)

    # Handle code files
    if file_extension in EXTENSION_TO_LANGUAGE:
        return extract_text_from_code(file_path)

    # Handle document files
    extraction_func = EXTRACTION_FUNCTIONS.get(file_extension)
    if extraction_func:
        try:
            return extraction_func(file_path)
        except TextExtractionError as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    else:
        logger.error(f"Unsupported file type: {file_path}")
        return ""
